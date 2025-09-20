from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
import os, tempfile, shutil, time, json, hashlib, re

# Excel / PDF
from openpyxl import load_workbook
from PyPDF2 import PdfReader

# OCR
from pdf2image import convert_from_path
import pytesseract

load_dotenv()

app = FastAPI(title=os.getenv("APP_TITLE", "自動読み取り機"))
security = HTTPBasic()

ADMIN_USER = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASS = os.getenv("ADMIN_PASSWORD", "password")
STORAGE_DIR = Path(os.getenv("STORAGE_DIR", "./data/work"))
RETENTION_MIN = int(os.getenv("RETENTION_MINUTES", "60"))
TEMPLATE_DIR = Path(os.getenv("TEMPLATE_DIR", "./data/templates"))
MAPPING_FILE = Path(os.getenv("MAPPING_FILE", "./data/config/mapping.json"))

STORAGE_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

def auth(creds: HTTPBasicCredentials = Depends(security)):
    if not (creds.username == ADMIN_USER and creds.password == ADMIN_PASS):
        raise HTTPException(status_code=401, detail="Unauthorized")
    return True

@app.get("/", response_class=HTMLResponse)
def index(_: bool = Depends(auth)):
    return """
    <html><body>
      <h1>自動読み取り機</h1>
      <p>Excelは「A列=項目名 / B列=値」、PDFは「項目名：値」の行を想定。</p>
      <form action="/process" method="post" enctype="multipart/form-data">
        <p>ファイルを選択: <input type="file" name="file" required></p>
        <p><button type="submit">アップロードして処理</button></p>
      </form>
      <hr/>
      <form action="/gc" method="post"><button type="submit">古いデータを削除(GC)</button></form>
    </body></html>
    """

@app.get("/healthz")
def healthz():
    return {"ok": True}

# ========== 差し込みロジック ==========
_ZWS = "\u200b"
def _normalize_text(s: str) -> str:
    return s.replace("｛", "{").replace("｝", "}").replace(_ZWS, "")

def _replace_text_block(raw_text: str, mapping: dict) -> str:
    text = _normalize_text(raw_text)
    for k, v in mapping.items():
        pattern = re.compile(r"\{\{\s*" + re.escape(k) + r"\s*\}\}")
        text = pattern.sub("" if v is None else str(v), text)
    return text

def _replace_in_paragraph(paragraph, mapping: dict):
    if not paragraph.runs:
        return
    joined = "".join(r.text for r in paragraph.runs)
    replaced = _replace_text_block(joined, mapping)
    if replaced != joined:
        paragraph.runs[0].text = replaced
        for r in paragraph.runs[1:]:
            r.text = ""

def _replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)

def fill_docx_placeholders(template_path: Path, output_path: Path, mapping: dict):
    doc = Document(str(template_path))
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in section.header.tables:
                _replace_in_table(t, mapping)
        if section.footer:
            for p in section.footer.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in section.footer.tables:
                _replace_in_table(t, mapping)
    doc.save(str(output_path))
# ===================================

# ====== テンプレ用のキー ======
PLACEHOLDER_KEYS = [
    "候補者氏名","フリガナ","生年月日","年齢","最終学歴","電話番号","MAIL","郵便番号",
    "現住所都道府県名","現年収","希望年収","就業可能時期","現職の雇用形態","ステータス",
    "運転免許の有無","自家用車の有無","既往歴の有無","扶養_配偶者の有無",
    "外国籍_永住権の有無","犯罪歴の有無","入社後2ヶ月までの費用","希望勤務地",
    "出張_可否","出張_期間","出張_範囲","転居_可否","転居_範囲",
    "紹介先企業名_1","紹介先企業名_2","紹介先企業名_3",
    "希望職種_1","希望職種_2","希望職種_3",
    "希望選考方法","面接候補日_1","面接候補日_2","面接候補日_3",
    "空白期間","中途退学理由","短期離職理由","推薦文"
]

ALIASES = {
    "扶養＆配偶者の有無": "扶養_配偶者の有無",
    "現住所（都道府県名）": "現住所都道府県名",
    "永住権の有無": "外国籍_永住権の有無",
    "入社後〜2ヶ月の費用": "入社後2ヶ月までの費用",
}

def normalize_label(label: str) -> str:
    if label is None: return ""
    s = str(label).strip().replace("：", ":").replace("　", " ")
    if s in ALIASES: return ALIASES[s]
    return s

def to_template_key(label: str) -> str:
    lab = normalize_label(label)
    if lab in PLACEHOLDER_KEYS: return lab
    for k in PLACEHOLDER_KEYS:
        if lab.startswith(k): return k
    return lab

# ====== Excel ======
def extract_from_excel(xlsx_path: Path) -> dict:
    wb = load_workbook(filename=str(xlsx_path), data_only=True)
    ws = wb.worksheets[0]
    result = {}
    for row in ws.iter_rows(min_row=1, values_only=True):
        if not row: continue
        key, val = (row[0], row[1] if len(row) > 1 else None)
        if key and val:
            tkey = to_template_key(str(key))
            if tkey in PLACEHOLDER_KEYS:
                result[tkey] = str(val).strip()
    return result

# ====== PDF（OCR fallbackあり） ======
def extract_from_pdf(pdf_path: Path) -> dict:
    text = ""
    try:
        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            text += page.extract_text() or ""
            text += "\n"
    except Exception:
        pass

    if not text.strip():
        # OCR fallback
        images = convert_from_path(str(pdf_path), dpi=300)
        ocr_texts = [pytesseract.image_to_string(img, lang="jpn") for img in images]
        text = "\n".join(ocr_texts)

    result = {}
    for line in text.splitlines():
        if not line.strip(): continue
        m = re.match(r"^\s*(.+?)\s*[:：]\s*(.+?)\s*$", line)
        if not m: continue
        raw_k, raw_v = m.group(1), m.group(2)
        tkey = to_template_key(raw_k)
        if tkey in PLACEHOLDER_KEYS and raw_v:
            result[tkey] = raw_v.strip()
    return result

# ===== プレビュー =====
def safe_preview_table(data: dict) -> str:
    def esc(s): return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    rows = "".join(f"<tr><th>{esc(k)}</th><td>{esc(v)}</td></tr>" for k,v in data.items())
    return f"<table border='1'>{rows}</table>"

# ===== メイン処理 =====
@app.post("/process")
async def process(_: bool = Depends(auth), file: UploadFile = File(...)):
    tmpdir = Path(tempfile.mkdtemp(dir=STORAGE_DIR))
    suffix = Path(file.filename).suffix.lower()
    inpath = tmpdir / f"input{suffix}"
    with open(inpath, "wb") as f:
        shutil.copyfileobj(file.file, f)

    extracted: dict = {}
    if suffix in [".xlsx", ".xlsm"]:
        extracted = extract_from_excel(inpath)
    elif suffix == ".pdf":
        extracted = extract_from_pdf(inpath)

    completed = {k: extracted.get(k, "") for k in PLACEHOLDER_KEYS}

    template_path = TEMPLATE_DIR / "template.docx"
    if not template_path.exists():
        raise HTTPException(500, f"テンプレートが見つかりません: {template_path}")

    out_docx = tmpdir / "output.docx"
    fill_docx_placeholders(template_path, out_docx, completed)

    meta = {
        "timestamp": int(time.time()),
        "input_hash": hashlib.sha256(inpath.read_bytes()).hexdigest()[:12],
        "filled_count": sum(1 for v in completed.values() if v),
        "missing_keys": [k for k, v in completed.items() if not v],
    }
    with open(tmpdir / "report.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    download_url = f"/download/{tmpdir.name}/output.docx"
    report_url   = f"/download/{tmpdir.name}/report.json"
    html = f"""
    <html><body>
      <h2>差し込み完了</h2>
      <p><a href="{download_url}">▶ Wordファイルをダウンロード</a></p>
      <p><a href="{report_url}">▶ レポートを表示</a></p>
      <h3>抽出プレビュー</h3>
      {safe_preview_table({k:v for k,v in completed.items() if v})}
    </body></html>
    """
    return HTMLResponse(content=html)

@app.get("/download/{job}/{fname}")
def download(job: str, fname: str, _: bool = Depends(auth)):
    path = STORAGE_DIR / job / fname
    if not path.exists(): raise HTTPException(404)
    return FileResponse(path)

@app.post("/gc")
def gc(_: bool = Depends(auth)):
    now = time.time()
    removed = 0
    for p in STORAGE_DIR.iterdir():
        if p.is_dir() and now - p.stat().st_mtime > RETENTION_MIN * 60:
            shutil.rmtree(p, ignore_errors=True); removed += 1
    return {"removed_dirs": removed}
